"""
LCEL 기반 RAG 체인 모듈.

LangChain Expression Language(LCEL)의 파이프 연산자(|)를 사용하여
Retriever → Prompt → LLM → OutputParser 순서로 체인을 조립한다.
ChromaDB가 없으면 data/docs/ 원본 문서를 파싱하여 자동 구축한다.

사용 방법:
    chain, retriever = get_rag_chain()
    answer = chain.invoke({"question": "병가 신청 시 증빙 서류가 필요한가요?", "history": "없음"})
"""

import os
from operator import itemgetter
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from langchain.schema import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate

load_dotenv()

# =============================================================
# 출처 강제 + "모르면 확인되지 않음" RAG 프롬프트 템플릿
# =============================================================
RAG_SYSTEM_PROMPT = """당신은 메타코딩 사내 문서 Q&A 비서입니다.
아래에 제공된 문서(Context)만 사용하여 질문에 답변하십시오.

규칙:
1. 반드시 제공된 문서에서만 근거를 찾아 답변하시오.
2. 문서에서 답을 찾을 수 없으면 "해당 내용은 제공된 문서에서 확인되지 않습니다."라고 답하시오.
3. 답변 마지막에 근거 문서명을 반드시 명시하시오. 형식: [출처: 문서명]
4. 추측이나 외부 지식을 사용하지 마시오.

Context (제공된 문서):
{context}

이전 대화:
{history}
"""

RAG_HUMAN_PROMPT = "질문: {question}"


def _build_llm() -> Any:
    """
    .env의 LLM_PROVIDER 값에 따라 LLM 인스턴스를 생성하여 반환한다.

    Returns:
        LangChain LLM 인스턴스 (ChatOllama | ChatOpenAI)

    Raises:
        ValueError: 지원하지 않는 LLM_PROVIDER 값일 때
    """
    provider = os.getenv("LLM_PROVIDER", "ollama").lower()

    if provider == "ollama":
        from langchain_ollama import ChatOllama
        return ChatOllama(
            base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),
            model=os.getenv("OLLAMA_MODEL", "deepseek-r1:8b"),
            temperature=0.1,
        )
    elif provider == "openai":
        from langchain_openai import ChatOpenAI
        return ChatOpenAI(
            api_key=os.getenv("OPENAI_API_KEY"),
            model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
            temperature=0.1,
        )
    else:
        raise ValueError(
            f"지원하지 않는 LLM_PROVIDER입니다: '{provider}'. "
            "ollama 또는 openai 중 하나를 선택하세요."
        )


def _build_retriever() -> Any:
    """
    ChromaDB Retriever를 생성한다.
    ChromaDB가 없으면 data/docs/ 원본 문서를 파싱하여 자동 구축한다.

    Returns:
        LangChain Retriever 인스턴스
    """
    # === INPUT ===
    chroma_dir = os.getenv("CHROMA_PERSIST_DIR", "./data/chroma_db")
    collection_name = os.getenv("CHROMA_COLLECTION_NAME", "metacoding_documents")
    top_k = int(os.getenv("RETRIEVER_TOP_K", "5"))
    embedding_model_name = os.getenv("EMBEDDING_MODEL", "jhgan/ko-sroberta-multitask")

    # === PROCESS ===
    from langchain_chroma import Chroma
    from langchain_community.embeddings import HuggingFaceEmbeddings

    embeddings = HuggingFaceEmbeddings(
        model_name=embedding_model_name,
        model_kwargs={"device": "cpu"},
    )

    # ChromaDB 데이터 존재 여부 확인
    has_data = os.path.isdir(chroma_dir) and any(
        f.endswith(".sqlite3") for f in _list_dir_recursive(chroma_dir)
    )

    if has_data:
        print(f"[INFO] 기존 ChromaDB 로드: {chroma_dir}")
        vectorstore = Chroma(
            collection_name=collection_name,
            embedding_function=embeddings,
            persist_directory=chroma_dir,
        )
    else:
        # ChromaDB 없음 → data/docs/ 원본 문서에서 자동 구축
        print("[INFO] ChromaDB가 없습니다. data/docs/ 원본 문서에서 자동 구축합니다.")
        docs = _parse_and_chunk_docs()
        if not docs:
            raise RuntimeError(
                "data/docs/ 디렉토리에 문서가 없습니다. "
                "원본 문서(PDF/DOCX/XLSX)를 data/docs/에 넣고 다시 실행하십시오."
            )
        vectorstore = Chroma.from_documents(
            documents=docs,
            embedding=embeddings,
            persist_directory=chroma_dir,
            collection_name=collection_name,
        )
        print(f"[INFO] ChromaDB 자동 구축 완료: {len(docs)}건 → {chroma_dir}")

    # === OUTPUT ===
    return vectorstore.as_retriever(search_kwargs={"k": top_k})


def _list_dir_recursive(path: str) -> list[str]:
    """
    디렉토리 내 모든 파일 이름을 재귀적으로 반환한다.

    Args:
        path: 탐색할 디렉토리 경로

    Returns:
        파일명 목록
    """
    result = []
    for root, _, files in os.walk(path):
        result.extend(files)
    return result


def _parse_and_chunk_docs(
    chunk_size: int = 500,
    overlap: int = 100,
) -> list[Document]:
    """data/docs/의 원본 PDF/DOCX/XLSX를 파싱→청킹하여 Document 리스트로 반환한다."""
    import pypdf
    from docx import Document as DocxDocument
    import openpyxl

    docs_dir = Path(__file__).parent.parent / "data" / "docs"
    if not docs_dir.exists():
        return []

    documents: list[Document] = []
    for file_path in sorted(docs_dir.rglob("*")):
        suffix = file_path.suffix.lower()
        source = file_path.stem

        texts: list[tuple[str, int]] = []  # (text, page_num)

        if suffix == ".pdf":
            try:
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    for page_num, page in enumerate(reader.pages, start=1):
                        text = (page.extract_text() or "").strip()
                        if text:
                            texts.append((text, page_num))
            except Exception:
                continue

        elif suffix == ".docx":
            try:
                doc = DocxDocument(str(file_path))
                full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if full_text:
                    texts.append((full_text, 1))
            except Exception:
                continue

        elif suffix == ".xlsx":
            try:
                wb = openpyxl.load_workbook(str(file_path), data_only=True)
                for idx, name in enumerate(wb.sheetnames, start=1):
                    ws = wb[name]
                    rows = []
                    for row in ws.iter_rows():
                        cells = [str(c.value).strip() for c in row if c.value is not None]
                        if cells:
                            rows.append(" | ".join(cells))
                    if rows:
                        texts.append(("\n".join(rows), idx))
            except Exception:
                continue

        # 청킹: 500자 단위, 100자 오버랩
        for text, page_num in texts:
            step = chunk_size - overlap
            start = 0
            while start < len(text):
                chunk = text[start:start + chunk_size].strip()
                if chunk:
                    documents.append(Document(
                        page_content=chunk,
                        metadata={"source": source, "page": page_num},
                    ))
                start += step

    return documents


def _format_docs(docs: list[Document]) -> str:
    """
    검색된 Document 목록을 프롬프트에 삽입할 텍스트 형식으로 변환한다.

    Args:
        docs: LangChain Document 객체 목록

    Returns:
        포맷팅된 컨텍스트 문자열
    """
    parts = []
    for i, doc in enumerate(docs, start=1):
        source = doc.metadata.get("source", "알 수 없음")
        page = doc.metadata.get("page", "-")
        parts.append(f"[문서 {i}] 출처: {source} (p.{page})\n{doc.page_content}")
    return "\n\n".join(parts)


def build_rag_chain() -> tuple[Any, Any]:
    """
    LCEL 파이프 연산자(|)로 RAG 체인과 Retriever를 조립하여 반환한다.

    체인 입력 형식:
        {"question": "사용자 질문", "history": "이전 대화 텍스트"}

    체인 구조:
        itemgetter("question") → Retriever → _format_docs  →┐
        itemgetter("history")  ─────────────────────────────┤→ Prompt → LLM → StrOutputParser
        itemgetter("question") ─────────────────────────────┘

    Returns:
        (chain, retriever) 튜플

    Raises:
        ValueError: LLM_PROVIDER 설정 오류 시
    """
    # === INPUT ===
    llm = _build_llm()                # ① LLM 인스턴스 생성
    retriever = _build_retriever()    # ② Retriever 생성 (ChromaDB)

    # === PROCESS ===
    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", RAG_SYSTEM_PROMPT),
            ("human", RAG_HUMAN_PROMPT),
        ]
    )

    # LCEL 파이프: 입력 dict에서 각 키를 꺼내 병렬 처리 후 프롬프트로 합침
    # itemgetter를 사용해 question 키만 추출하여 retriever로 전달
    chain = (
        {
            "context": itemgetter("question") | retriever | _format_docs,  # ③ 질문으로 문서 검색 후 포맷
            "history": itemgetter("history"),                               # ④ 이전 대화 히스토리 추출
            "question": itemgetter("question"),                             # ⑤ 질문 그대로 전달
        }
        | prompt              # ⑥ 시스템+사용자 프롬프트 조합
        | llm                 # ⑦ LLM 호출
        | StrOutputParser()   # ⑧ 문자열로 파싱
    )

    # === OUTPUT ===
    return chain, retriever


# 싱글턴 캐시 (앱 시작 시 1회만 초기화)
_rag_chain_cache: Any = None
_retriever_cache: Any = None


def get_rag_chain() -> tuple[Any, Any]:
    """
    RAG 체인과 Retriever 싱글턴 인스턴스를 반환한다.
    최초 호출 시 초기화하고 이후 캐시된 값을 재사용한다.

    Returns:
        (chain, retriever) 튜플
    """
    global _rag_chain_cache, _retriever_cache
    if _rag_chain_cache is None:
        _rag_chain_cache, _retriever_cache = build_rag_chain()
    return _rag_chain_cache, _retriever_cache
