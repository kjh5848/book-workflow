"""샘플 문서 생성 스크립트.

비교 실험을 위한 샘플 PDF, DOCX, XLSX 문서를 생성한다.

IPO 패턴:
  Input  - 없음 (내장 데이터)
  Process - reportlab(PDF), python-docx(DOCX), openpyxl(XLSX)로 문서 생성
  Output - data/ 폴더에 3개 샘플 문서

실행 방법:
    python data/create_sample_docs.py
"""

import os
import sys
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
CAPTURED_DIR = BASE_DIR / "captured"


# ============================================================
# 1. 샘플 PDF 생성 (취업규칙 + 연봉 테이블 + 조직도 이미지)
# ============================================================

def create_sample_pdf(output_path: Path) -> None:
    """취업규칙 샘플 PDF를 생성한다.

    Args:
        output_path: 저장할 PDF 파일 경로.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        print("reportlab 패키지가 필요합니다: pip install reportlab")
        return

    # --- 한글 폰트 등록 ---
    font_registered = False
    font_candidates = [
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "C:/Windows/Fonts/malgun.ttf",
    ]
    font_name = "Helvetica"

    for font_path in font_candidates:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("KoreanFont", font_path))
                font_name = "KoreanFont"
                font_registered = True
                break
            except Exception:
                continue

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleKo", parent=styles["Title"],
        fontName=font_name, fontSize=18, spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "HeadingKo", parent=styles["Heading2"],
        fontName=font_name, fontSize=14, spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BodyKo", parent=styles["Normal"],
        fontName=font_name, fontSize=10, leading=14, spaceAfter=6,
    )

    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    elements = []

    # 표지
    elements.append(Spacer(1, 40 * mm))
    elements.append(Paragraph("MetaCoding Inc.", title_style))
    elements.append(Paragraph("Employment Rules v1.0", heading_style))
    elements.append(Spacer(1, 10 * mm))
    elements.append(Paragraph("2025.01.01", body_style))
    elements.append(Spacer(1, 30 * mm))

    # 본문
    sections = [
        ("1. General Provisions", [
            "Article 1 (Purpose): These rules prescribe the working conditions and duties of employees at MetaCoding Inc.",
            "Article 2 (Scope): These rules apply to all full-time employees of the company.",
            "Article 3 (Definitions): 'Employee' refers to a person who has signed an employment contract with the company.",
        ]),
        ("2. Working Hours and Leave", [
            "Article 10 (Working Hours): Standard working hours are 8 hours per day, 40 hours per week.",
            "Article 11 (Break Time): Employees are entitled to a 1-hour lunch break between 12:00 and 13:00.",
            "Article 12 (Annual Leave): Employees with 1 year of service receive 15 days of annual leave.",
            "Article 13 (Sick Leave): Employees may use up to 30 days of paid sick leave per year with a doctor's note.",
        ]),
        ("3. Salary Structure", [
            "Article 20 (Salary Payment): Salaries are paid on the 25th of each month.",
            "Article 21 (Salary Components): Salary consists of base pay, position allowance, and performance bonus.",
        ]),
    ]

    for title, paragraphs in sections:
        elements.append(Paragraph(title, heading_style))
        for p in paragraphs:
            elements.append(Paragraph(p, body_style))
        elements.append(Spacer(1, 5 * mm))

    # 연봉 테이블
    elements.append(Paragraph("Salary Table by Position", heading_style))
    salary_data = [
        ["Position", "Base Salary", "Allowance", "Total"],
        ["Staff", "35,000,000", "3,000,000", "38,000,000"],
        ["Senior", "45,000,000", "5,000,000", "50,000,000"],
        ["Manager", "55,000,000", "7,000,000", "62,000,000"],
        ["Director", "70,000,000", "10,000,000", "80,000,000"],
        ["VP", "90,000,000", "15,000,000", "105,000,000"],
    ]
    salary_table = Table(salary_data, colWidths=[80, 100, 80, 100])
    salary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
    ]))
    elements.append(salary_table)
    elements.append(Spacer(1, 10 * mm))

    # 조직도 (간단 테이블로 표현)
    elements.append(Paragraph("Organization Chart", heading_style))
    org_data = [
        ["", "", "CEO", "", ""],
        ["", "", "|", "", ""],
        ["CTO", "---", "COO", "---", "CFO"],
        ["|", "", "|", "", "|"],
        ["Dev Team", "", "Operations", "", "Finance"],
    ]
    org_table = Table(org_data, colWidths=[80, 30, 80, 30, 80])
    org_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("BOX", (2, 0), (2, 0), 1, colors.black),
        ("BOX", (0, 2), (0, 2), 1, colors.black),
        ("BOX", (2, 2), (2, 2), 1, colors.black),
        ("BOX", (4, 2), (4, 2), 1, colors.black),
    ]))
    elements.append(org_table)

    doc.build(elements)
    print(f"  PDF 생성 완료: {output_path.name}")


# ============================================================
# 2. 샘플 DOCX 생성 (매출 보고서 + 표 + 임베디드 이미지)
# ============================================================

def create_sample_docx(output_path: Path) -> None:
    """매출 보고서 샘플 DOCX를 생성한다.

    Args:
        output_path: 저장할 DOCX 파일 경로.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("python-docx 패키지가 필요합니다: pip install python-docx")
        return

    doc = Document()

    # 스타일 설정
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    # 제목
    title = doc.add_heading("2025 H1 Sales Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("MetaCoding Inc. | Confidential")
    doc.add_paragraph("Report Date: 2025-07-01")
    doc.add_paragraph("")

    # 요약
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Total revenue for H1 2025 reached KRW 12.5 billion, "
        "representing a 15% increase compared to H1 2024. "
        "The Sales department led with KRW 5.2 billion in revenue."
    )

    # 매출 테이블
    doc.add_heading("2. Revenue by Department", level=1)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Department", "Q1 (KRW)", "Q2 (KRW)", "Total (KRW)"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    rows_data = [
        ["Sales", "2,500,000,000", "2,700,000,000", "5,200,000,000"],
        ["Development", "1,800,000,000", "2,000,000,000", "3,800,000,000"],
        ["Marketing", "800,000,000", "900,000,000", "1,700,000,000"],
        ["Operations", "700,000,000", "600,000,000", "1,300,000,000"],
        ["Total", "5,800,000,000", "6,200,000,000", "12,000,000,000"],
    ]
    for row_idx, row_data in enumerate(rows_data, 1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    doc.add_paragraph("")

    # 차트 이미지 대체 (matplotlib로 생성)
    doc.add_heading("3. Revenue Trend Chart", level=1)
    chart_path = _create_chart_image(output_path.parent)
    if chart_path and chart_path.exists():
        doc.add_picture(str(chart_path), width=Inches(5))
        chart_path.unlink(missing_ok=True)
    else:
        doc.add_paragraph("[Chart: Revenue trend visualization]")

    # 전략 섹션
    doc.add_heading("4. H2 Strategy", level=1)
    doc.add_paragraph(
        "Key strategies for H2 2025 include:"
    )
    strategies = [
        "Expand enterprise client base by 20%",
        "Launch new SaaS product line",
        "Invest in AI/ML capabilities for product enhancement",
        "Strengthen partnership program with 10 new partners",
    ]
    for s in strategies:
        doc.add_paragraph(s, style="List Bullet")

    doc.save(str(output_path))
    print(f"  DOCX 생성 완료: {output_path.name}")


def _create_chart_image(parent_dir: Path) -> Path | None:
    """매출 트렌드 차트 이미지를 생성한다."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun"]
        sales = [850, 920, 780, 1050, 1100, 1150]
        dev = [600, 580, 620, 650, 680, 720]

        fig, ax = plt.subplots(figsize=(8, 4))
        ax.plot(months, sales, "o-", label="Sales", color="#d4af37", linewidth=2)
        ax.plot(months, dev, "s-", label="Development", color="#333333", linewidth=2)
        ax.set_title("Monthly Revenue Trend (H1 2025)", fontsize=13)
        ax.set_ylabel("Revenue (M KRW)")
        ax.legend()
        ax.grid(True, alpha=0.3)
        fig.tight_layout()

        chart_path = parent_dir / "_temp_chart.png"
        fig.savefig(str(chart_path), dpi=150)
        plt.close(fig)
        return chart_path

    except ImportError:
        return None


# ============================================================
# 3. 샘플 XLSX 생성 (부서별 예산 + 차트 시트)
# ============================================================

def create_sample_xlsx(output_path: Path) -> None:
    """부서별 예산 샘플 XLSX를 생성한다.

    Args:
        output_path: 저장할 XLSX 파일 경로.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.chart import BarChart, Reference
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("openpyxl 패키지가 필요합니다: pip install openpyxl")
        return

    wb = Workbook()

    # --- Sheet 1: 부서별 예산 ---
    ws1 = wb.active
    ws1.title = "Budget Summary"

    header_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    headers = ["Department", "Personnel", "Equipment", "Marketing", "R&D", "Total"]
    for col, header in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        cell.border = border
        ws1.column_dimensions[get_column_letter(col)].width = 15

    budget_data = [
        ["Sales", 1200, 300, 800, 0, 2300],
        ["Development", 2500, 1500, 0, 2000, 6000],
        ["Marketing", 800, 200, 3000, 0, 4000],
        ["Operations", 1500, 800, 0, 0, 2300],
        ["HR", 600, 100, 0, 0, 700],
    ]

    alt_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
    for row_idx, row_data in enumerate(budget_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws1.cell(row=row_idx, column=col_idx, value=value)
            cell.border = border
            if col_idx > 1:
                cell.number_format = "#,##0"
                cell.alignment = Alignment(horizontal="right")
            if row_idx % 2 == 0:
                cell.fill = alt_fill

    # 합계 행
    total_row = len(budget_data) + 2
    ws1.cell(row=total_row, column=1, value="Total").font = Font(bold=True)
    for col in range(2, 7):
        cell = ws1.cell(
            row=total_row, column=col,
            value=sum(row[col - 1] for row in budget_data),
        )
        cell.font = Font(bold=True)
        cell.number_format = "#,##0"
        cell.border = border

    # --- Sheet 2: 차트 ---
    ws2 = wb.create_sheet("Budget Chart")

    # 차트 데이터 복사
    for row_idx, row_data in enumerate(budget_data, 2):
        ws2.cell(row=row_idx, column=1, value=row_data[0])
        ws2.cell(row=row_idx, column=2, value=row_data[5])  # Total

    ws2.cell(row=1, column=1, value="Department")
    ws2.cell(row=1, column=2, value="Total Budget")

    chart = BarChart()
    chart.type = "col"
    chart.title = "Department Budget (Unit: M KRW)"
    chart.y_axis.title = "Budget (M KRW)"
    chart.x_axis.title = "Department"

    data = Reference(ws2, min_col=2, min_row=1, max_row=len(budget_data) + 1)
    cats = Reference(ws2, min_col=1, min_row=2, max_row=len(budget_data) + 1)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.shape = 4
    ws2.add_chart(chart, "D2")

    # --- Sheet 3: 월별 실적 ---
    ws3 = wb.create_sheet("Monthly Actuals")
    monthly_headers = ["Month", "Sales", "Development", "Marketing", "Operations"]
    for col, header in enumerate(monthly_headers, 1):
        cell = ws3.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = border

    monthly_data = [
        ["Jan", 420, 510, 350, 200],
        ["Feb", 380, 490, 320, 190],
        ["Mar", 450, 530, 380, 210],
        ["Apr", 510, 560, 410, 230],
        ["May", 480, 540, 390, 220],
        ["Jun", 520, 580, 420, 250],
    ]
    for row_idx, row_data in enumerate(monthly_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws3.cell(row=row_idx, column=col_idx, value=value)
            cell.border = border
            if col_idx > 1:
                cell.number_format = "#,##0"

    wb.save(str(output_path))
    print(f"  XLSX 생성 완료: {output_path.name}")


# ============================================================
# 메인 실행
# ============================================================

def main() -> None:
    """3종 샘플 문서를 생성한다."""
    print("=" * 50)
    print("샘플 문서 생성 시작")
    print("=" * 50)

    # captured 디렉토리 생성
    for subdir in ["pdf", "docx", "xlsx"]:
        (CAPTURED_DIR / subdir).mkdir(parents=True, exist_ok=True)

    # PDF 생성
    pdf_path = BASE_DIR / "sample_hr_policy.pdf"
    create_sample_pdf(pdf_path)

    # DOCX 생성
    docx_path = BASE_DIR / "sample_sales_report.docx"
    create_sample_docx(docx_path)

    # XLSX 생성
    xlsx_path = BASE_DIR / "sample_budget.xlsx"
    create_sample_xlsx(xlsx_path)

    print("\n" + "=" * 50)
    print("샘플 문서 생성 완료")
    print(f"  - {pdf_path.name}")
    print(f"  - {docx_path.name}")
    print(f"  - {xlsx_path.name}")
    print(f"  - captured/ 디렉토리 준비 완료")
    print("=" * 50)


if __name__ == "__main__":
    main()
