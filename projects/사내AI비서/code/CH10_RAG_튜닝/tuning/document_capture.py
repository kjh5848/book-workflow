"""문서 캡처 + 인제스천 파이프라인 모듈.

문서를 페이지/시트별 PNG로 캡처하고, 텍스트+메타데이터+이미지 참조를
벡터DB에 저장하는 인제스천 파이프라인을 구현한다.

IPO 패턴:
  Input  - PDF/DOCX/XLSX 문서 파일
  Process - 캡처(PNG) + 텍스트 추출 + 메타데이터 + 벡터DB 저장
  Output - data/captured/ 이미지 + ChromaDB 인제스천

실행 방법:
    python tuning/document_capture.py
"""

import os
import sys
import json
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from rich.console import Console
from rich.table import Table

load_dotenv()

console = Console()

BASE_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = BASE_DIR / "data"
CAPTURED_DIR = DATA_DIR / "captured"
OUTPUTS_DIR = BASE_DIR / "outputs"


# ============================================================
# 1. PDF 페이지별 캡처
# ============================================================

def capture_pdf_pages(pdf_path: Path) -> list[dict[str, Any]]:
    """PDF를 페이지별 PNG로 캡처하고 텍스트를 추출한다.

    Args:
        pdf_path: PDF 파일 경로.

    Returns:
        페이지별 캡처 결과 리스트.
    """
    results = []
    output_dir = CAPTURED_DIR / "pdf"
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(pdf_path))
        stem = pdf_path.stem

        for page_num in range(len(doc)):
            page = doc[page_num]

            # 페이지를 PNG로 렌더링
            pix = page.get_pixmap(dpi=200)
            img_filename = f"{stem}_page_{page_num + 1}.png"
            img_path = output_dir / img_filename
            pix.save(str(img_path))

            # 텍스트 추출
            text = page.get_text()

            results.append({
                "page": page_num + 1,
                "image_path": str(img_path),
                "text": text,
                "metadata": {
                    "source": pdf_path.name,
                    "page": page_num + 1,
                    "total_pages": len(doc),
                    "format": "pdf",
                    "image_path": str(img_path),
                },
            })

        doc.close()
        console.print(f"  [green]PDF 캡처 완료: {len(results)}페이지 → {output_dir}[/green]")

    except ImportError:
        console.print("[yellow]PyMuPDF 없음. pip install PyMuPDF 를 실행하십시오.[/yellow]")

        # 대체: pypdf로 텍스트만 추출
        try:
            import pypdf

            reader = pypdf.PdfReader(str(pdf_path))
            for page_num, page in enumerate(reader.pages):
                results.append({
                    "page": page_num + 1,
                    "image_path": "",
                    "text": page.extract_text() or "",
                    "metadata": {
                        "source": pdf_path.name,
                        "page": page_num + 1,
                        "format": "pdf",
                    },
                })
        except ImportError:
            console.print("[red]pypdf도 없습니다. pip install pypdf 를 실행하십시오.[/red]")

    return results


# ============================================================
# 2. DOCX 임베디드 이미지 추출
# ============================================================

def capture_docx_images(docx_path: Path) -> list[dict[str, Any]]:
    """DOCX에서 임베디드 이미지를 추출하고 본문 텍스트를 반환한다.

    Args:
        docx_path: DOCX 파일 경로.

    Returns:
        이미지별 캡처 결과 리스트.
    """
    results = []
    output_dir = CAPTURED_DIR / "docx"
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document

        doc = Document(str(docx_path))
        stem = docx_path.stem

        # 본문 텍스트 추출
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # 임베디드 이미지 추출
        img_idx = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_idx += 1
                ext = rel.target_ref.split(".")[-1] if "." in rel.target_ref else "png"
                img_filename = f"{stem}_img_{img_idx}.{ext}"
                img_path = output_dir / img_filename

                try:
                    with open(img_path, "wb") as f:
                        f.write(rel.target_part.blob)

                    results.append({
                        "image_path": str(img_path),
                        "alt_text": f"Image {img_idx} from {docx_path.name}",
                        "context_text": full_text[:200],
                        "metadata": {
                            "source": docx_path.name,
                            "image_index": img_idx,
                            "format": "docx",
                            "image_path": str(img_path),
                        },
                    })
                except Exception as e:
                    console.print(f"  [yellow]이미지 {img_idx} 추출 실패: {e}[/yellow]")

        # 이미지가 없어도 텍스트는 결과에 포함
        if not results:
            results.append({
                "image_path": "",
                "alt_text": "",
                "context_text": full_text,
                "metadata": {
                    "source": docx_path.name,
                    "format": "docx",
                },
            })

        console.print(f"  [green]DOCX 이미지 추출 완료: {img_idx}개 → {output_dir}[/green]")

    except ImportError:
        console.print("[red]python-docx 패키지 필요: pip install python-docx[/red]")

    return results


# ============================================================
# 3. XLSX 시트별 테이블 이미지 렌더링
# ============================================================

def capture_xlsx_sheets(xlsx_path: Path) -> list[dict[str, Any]]:
    """XLSX 시트를 테이블 이미지로 렌더링하고 데이터를 추출한다.

    Args:
        xlsx_path: XLSX 파일 경로.

    Returns:
        시트별 캡처 결과 리스트.
    """
    results = []
    output_dir = CAPTURED_DIR / "xlsx"
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        from openpyxl import load_workbook

        wb = load_workbook(str(xlsx_path), data_only=True)
        stem = xlsx_path.stem

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # 데이터 텍스트 추출
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_str = " | ".join(str(v) if v is not None else "" for v in row)
                rows.append(row_str)
            data_text = "\n".join(rows)

            # 시트를 이미지로 렌더링
            img_path = _render_sheet_image(ws, stem, sheet_name, output_dir)

            results.append({
                "sheet_name": sheet_name,
                "image_path": str(img_path) if img_path else "",
                "data_text": data_text,
                "metadata": {
                    "source": xlsx_path.name,
                    "sheet": sheet_name,
                    "format": "xlsx",
                    "image_path": str(img_path) if img_path else "",
                    "row_count": ws.max_row or 0,
                    "col_count": ws.max_column or 0,
                },
            })

        console.print(f"  [green]XLSX 캡처 완료: {len(results)}시트 → {output_dir}[/green]")

    except ImportError:
        console.print("[red]openpyxl 패키지 필요: pip install openpyxl[/red]")

    return results


def _render_sheet_image(
    ws: Any,
    stem: str,
    sheet_name: str,
    output_dir: Path,
) -> Path | None:
    """워크시트를 matplotlib 테이블 이미지로 렌더링한다."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        data = []
        for row in ws.iter_rows(values_only=True):
            data.append([str(v) if v is not None else "" for v in row])

        if not data:
            return None

        fig, ax = plt.subplots(figsize=(12, max(3, len(data) * 0.4)))
        ax.axis("off")
        ax.set_title(f"{stem} - {sheet_name}", fontsize=12, pad=10)

        tbl = ax.table(
            cellText=data[1:] if len(data) > 1 else data,
            colLabels=data[0] if len(data) > 1 else None,
            cellLoc="center",
            loc="center",
        )
        tbl.auto_set_font_size(False)
        tbl.set_fontsize(8)
        tbl.scale(1.2, 1.5)

        # 헤더 스타일
        if len(data) > 1:
            for j in range(len(data[0])):
                tbl[0, j].set_facecolor("#333333")
                tbl[0, j].set_text_props(color="white", weight="bold")

        safe_name = sheet_name.replace(" ", "_")
        img_path = output_dir / f"{stem}_{safe_name}.png"
        fig.savefig(str(img_path), dpi=150, bbox_inches="tight")
        plt.close(fig)
        return img_path

    except ImportError:
        return None


# ============================================================
# 4. 메타데이터 추출
# ============================================================

def extract_metadata(file_path: Path) -> dict[str, Any]:
    """문서 파일에서 메타데이터를 추출한다.

    Args:
        file_path: 문서 파일 경로.

    Returns:
        메타데이터 딕셔너리.
    """
    stat = file_path.stat()
    metadata = {
        "filename": file_path.name,
        "format": file_path.suffix.lstrip("."),
        "size_bytes": stat.st_size,
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            info = reader.metadata
            if info:
                metadata["title"] = info.get("/Title", "")
                metadata["author"] = info.get("/Author", "")
                metadata["creator"] = info.get("/Creator", "")
            metadata["page_count"] = len(reader.pages)
        except (ImportError, Exception):
            pass

    elif suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(file_path))
            props = doc.core_properties
            metadata["title"] = props.title or ""
            metadata["author"] = props.author or ""
            metadata["created"] = props.created.isoformat() if props.created else ""
        except (ImportError, Exception):
            pass

    elif suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(file_path), read_only=True)
            metadata["sheet_names"] = wb.sheetnames
            metadata["sheet_count"] = len(wb.sheetnames)
            wb.close()
        except (ImportError, Exception):
            pass

    return metadata


# ============================================================
# 5. 벡터DB 인제스천
# ============================================================

def ingest_to_vectordb(
    documents: list[dict[str, Any]],
    collection_name: str = "document_captures",
) -> dict[str, Any]:
    """텍스트 청크와 이미지 참조를 ChromaDB에 저장한다.

    Args:
        documents: 캡처 결과 리스트 (text, metadata, image_path 포함).
        collection_name: ChromaDB 컬렉션 이름.

    Returns:
        인제스천 결과 요약.
    """
    try:
        import chromadb
        from chromadb.config import Settings

        client = chromadb.Client(Settings(anonymized_telemetry=False))
        collection = client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"},
        )

        ids = []
        texts = []
        metadatas = []

        for i, doc in enumerate(documents):
            text = doc.get("text") or doc.get("data_text") or doc.get("context_text") or ""
            if not text.strip():
                continue

            doc_id = f"doc_{i}_{hash(text[:50]) % 10000}"
            meta = doc.get("metadata", {})

            # 이미지 경로를 메타데이터에 포함
            image_path = doc.get("image_path", "")
            if image_path:
                meta["image_path"] = image_path

            ids.append(doc_id)
            texts.append(text[:5000])  # ChromaDB 제한
            metadatas.append({k: str(v) for k, v in meta.items()})

        if ids:
            collection.add(ids=ids, documents=texts, metadatas=metadatas)

        result = {
            "collection": collection_name,
            "total_documents": len(ids),
            "skipped": len(documents) - len(ids),
        }

        console.print(
            f"  [green]벡터DB 인제스천 완료: {len(ids)}건 저장[/green]"
        )
        return result

    except ImportError:
        console.print("[yellow]chromadb 패키지 필요: pip install chromadb[/yellow]")
        return {"error": "chromadb not installed", "total_documents": 0}


# ============================================================
# 데모 실행
# ============================================================

def run_capture_pipeline() -> None:
    """문서 캡처 + 인제스천 파이프라인 데모를 실행한다."""
    console.rule("[bold blue]CH10 문서 캡처 + 인제스천 파이프라인[/bold blue]")

    all_documents: list[dict[str, Any]] = []

    # --- PDF 캡처 ---
    pdf_files = list(DATA_DIR.glob("*.pdf")) + list((DATA_DIR / "docs").rglob("*.pdf"))
    if pdf_files:
        pdf_path = pdf_files[0]
        console.print(f"\n[bold]1. PDF 캡처: {pdf_path.name}[/bold]")

        meta = extract_metadata(pdf_path)
        console.print(f"   메타데이터: {json.dumps(meta, ensure_ascii=False, indent=2)[:200]}")

        pages = capture_pdf_pages(pdf_path)
        all_documents.extend(pages)

        # 결과 테이블
        table = Table(title="PDF 캡처 결과")
        table.add_column("페이지", style="cyan", justify="center")
        table.add_column("이미지", style="yellow")
        table.add_column("텍스트 길이", style="green", justify="right")

        for page in pages:
            img_name = Path(page["image_path"]).name if page["image_path"] else "(없음)"
            table.add_row(
                str(page["page"]),
                img_name,
                f"{len(page['text']):,}자",
            )
        console.print(table)
    else:
        console.print("\n[yellow]PDF 파일이 없습니다.[/yellow]")

    # --- DOCX 캡처 ---
    docx_files = list(DATA_DIR.glob("*.docx")) + list((DATA_DIR / "docs").rglob("*.docx"))
    if docx_files:
        docx_path = docx_files[0]
        console.print(f"\n[bold]2. DOCX 이미지 추출: {docx_path.name}[/bold]")
        images = capture_docx_images(docx_path)
        all_documents.extend(images)
        console.print(f"   추출된 항목: {len(images)}개")
    else:
        console.print("\n[yellow]DOCX 파일이 없습니다.[/yellow]")

    # --- XLSX 캡처 ---
    xlsx_files = list(DATA_DIR.glob("*.xlsx")) + list((DATA_DIR / "docs").rglob("*.xlsx"))
    if xlsx_files:
        xlsx_path = xlsx_files[0]
        console.print(f"\n[bold]3. XLSX 시트 캡처: {xlsx_path.name}[/bold]")
        sheets = capture_xlsx_sheets(xlsx_path)
        all_documents.extend(sheets)

        table = Table(title="XLSX 캡처 결과")
        table.add_column("시트", style="cyan")
        table.add_column("이미지", style="yellow")
        table.add_column("데이터 길이", style="green", justify="right")

        for sheet in sheets:
            img_name = Path(sheet["image_path"]).name if sheet["image_path"] else "(없음)"
            table.add_row(
                sheet["sheet_name"],
                img_name,
                f"{len(sheet.get('data_text', '')):,}자",
            )
        console.print(table)
    else:
        console.print("\n[yellow]XLSX 파일이 없습니다.[/yellow]")

    # --- 벡터DB 인제스천 ---
    if all_documents:
        console.print(f"\n[bold]4. 벡터DB 인제스천[/bold]")
        ingest_result = ingest_to_vectordb(all_documents)
        console.print(f"   결과: {json.dumps(ingest_result, ensure_ascii=False)}")

    # --- 요약 ---
    console.rule("[bold green]파이프라인 완료[/bold green]")

    summary_table = Table(title="캡처 파이프라인 요약")
    summary_table.add_column("항목", style="cyan")
    summary_table.add_column("결과", style="green")

    summary_table.add_row("총 문서", f"{len(all_documents)}건")
    summary_table.add_row("PDF 페이지", f"{len(pdf_files)}개 파일")
    summary_table.add_row("DOCX 이미지", f"{len(docx_files)}개 파일")
    summary_table.add_row("XLSX 시트", f"{len(xlsx_files)}개 파일")
    summary_table.add_row("캡처 경로", str(CAPTURED_DIR))

    console.print(summary_table)
    console.print(
        "\n[bold]인제스천 파이프라인 흐름:[/bold]\n"
        "  문서 → 페이지/시트 캡처(PNG) → 텍스트 추출\n"
        "  → 메타데이터 + 이미지 참조 포함 → ChromaDB 저장\n"
        "  → 검색 시 관련 텍스트 + 원본 이미지 경로 반환\n"
    )


if __name__ == "__main__":
    _PROJECT_ROOT = str(BASE_DIR)
    if _PROJECT_ROOT not in sys.path:
        sys.path.insert(0, _PROJECT_ROOT)

    run_capture_pipeline()
