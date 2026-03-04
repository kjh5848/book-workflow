"""문서 파싱 전략 비교 모듈.

라이브러리 파싱과 vLLM 파싱 두 전략을 같은 문서에 적용하여
속도, 텍스트 품질, 표/이미지 처리 능력을 비교한다.

IPO 패턴:
  Input  - PDF/DOCX/XLSX 문서 파일
  Process - 전략 1(라이브러리) vs 전략 2(vLLM/LLaVA)로 파싱
  Output - 비교표 (속도, 텍스트 길이, 표 추출 수, 이미지 설명 여부)

실행 방법:
    python tuning/document_parser.py
"""

import os
import sys
import time
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from rich.console import Console
from rich.table import Table

load_dotenv()

console = Console()

BASE_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = BASE_DIR / "data"
OUTPUTS_DIR = BASE_DIR / "outputs"
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
VISION_MODEL = os.getenv("VISION_MODEL", "llava:7b")


# ============================================================
# 전략 1: 라이브러리 파싱
# ============================================================

def parse_pdf_library(pdf_path: Path) -> dict[str, Any]:
    """라이브러리 기반 PDF 파싱 (pypdf + pdfplumber).

    Args:
        pdf_path: PDF 파일 경로.

    Returns:
        파싱 결과 딕셔너리.
    """
    result = {
        "strategy": "library",
        "format": "PDF",
        "text": "",
        "tables": [],
        "images": [],
        "elapsed": 0.0,
    }

    start = time.time()

    # pypdf로 텍스트 추출
    try:
        import pypdf

        reader = pypdf.PdfReader(str(pdf_path))
        pages_text = []
        image_count = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            pages_text.append(text)
            image_count += len(page.images)

        result["text"] = "\n\n".join(pages_text)
        result["images"] = [f"[image_{i+1}]" for i in range(image_count)]

    except ImportError:
        result["text"] = "[pypdf 패키지 필요: pip install pypdf]"

    # pdfplumber로 표 추출
    try:
        import pdfplumber

        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    result["tables"].append(table)

    except ImportError:
        pass  # pdfplumber 없으면 표 추출 생략
    except FileNotFoundError:
        pass  # 파일이 없으면 생략

    result["elapsed"] = time.time() - start
    return result


def parse_pdf_vllm(pdf_path: Path) -> dict[str, Any]:
    """vLLM(LLaVA) 기반 PDF 파싱.

    PDF를 페이지별 이미지로 변환 후 LLaVA로 구조화 MD를 생성한다.

    Args:
        pdf_path: PDF 파일 경로.

    Returns:
        파싱 결과 딕셔너리.
    """
    result = {
        "strategy": "vllm",
        "format": "PDF",
        "text": "",
        "tables": [],
        "images": [],
        "elapsed": 0.0,
    }

    start = time.time()

    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(pdf_path))
        page_texts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            pix = page.get_pixmap(dpi=150)

            img_path = OUTPUTS_DIR / f"_vllm_page_{page_num + 1}.png"
            img_path.parent.mkdir(parents=True, exist_ok=True)
            pix.save(str(img_path))

            caption = _call_llava(img_path)
            page_texts.append(f"## Page {page_num + 1}\n\n{caption}")
            result["images"].append(str(img_path))

            img_path.unlink(missing_ok=True)

        doc.close()
        result["text"] = "\n\n".join(page_texts)

        # LLaVA 출력에서 표 패턴 카운트
        table_markers = result["text"].count("|")
        if table_markers > 10:
            result["tables"].append("[LLaVA detected table structure]")

    except ImportError:
        result["text"] = "[PyMuPDF 패키지 필요: pip install PyMuPDF]"

    result["elapsed"] = time.time() - start
    return result


def parse_docx_library(docx_path: Path) -> dict[str, Any]:
    """라이브러리 기반 DOCX 파싱 (python-docx).

    Args:
        docx_path: DOCX 파일 경로.

    Returns:
        파싱 결과 딕셔너리.
    """
    result = {
        "strategy": "library",
        "format": "DOCX",
        "text": "",
        "tables": [],
        "images": [],
        "elapsed": 0.0,
    }

    start = time.time()

    try:
        from docx import Document

        doc = Document(str(docx_path))

        # 텍스트 추출
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        result["text"] = "\n".join(paragraphs)

        # 표 추출
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            result["tables"].append(table_data)

        # 임베디드 이미지 목록
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                result["images"].append(rel.target_ref)

    except ImportError:
        result["text"] = "[python-docx 패키지 필요: pip install python-docx]"

    result["elapsed"] = time.time() - start
    return result


def parse_docx_vllm(docx_path: Path) -> dict[str, Any]:
    """vLLM 기반 DOCX 파싱.

    DOCX 임베디드 이미지를 추출하고 LLaVA로 분석한다.

    Args:
        docx_path: DOCX 파일 경로.

    Returns:
        파싱 결과 딕셔너리.
    """
    result = {
        "strategy": "vllm",
        "format": "DOCX",
        "text": "",
        "tables": [],
        "images": [],
        "elapsed": 0.0,
    }

    start = time.time()

    try:
        from docx import Document

        doc = Document(str(docx_path))

        # 기본 텍스트 (python-docx에서)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        base_text = "\n".join(paragraphs)

        # 이미지 추출 + LLaVA 분석
        image_descriptions = []
        img_idx = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_idx += 1
                img_path = OUTPUTS_DIR / f"_vllm_docx_img_{img_idx}.png"
                img_path.parent.mkdir(parents=True, exist_ok=True)

                try:
                    with open(img_path, "wb") as f:
                        f.write(rel.target_part.blob)

                    caption = _call_llava(img_path)
                    image_descriptions.append(
                        f"### Image {img_idx}\n\n{caption}"
                    )
                    result["images"].append(str(img_path))
                    img_path.unlink(missing_ok=True)

                except Exception:
                    image_descriptions.append(
                        f"### Image {img_idx}\n\n[extraction failed]"
                    )

        result["text"] = base_text
        if image_descriptions:
            result["text"] += "\n\n" + "\n\n".join(image_descriptions)

    except ImportError:
        result["text"] = "[python-docx 패키지 필요: pip install python-docx]"

    result["elapsed"] = time.time() - start
    return result


def parse_xlsx_library(xlsx_path: Path) -> dict[str, Any]:
    """라이브러리 기반 XLSX 파싱 (openpyxl).

    Args:
        xlsx_path: XLSX 파일 경로.

    Returns:
        파싱 결과 딕셔너리.
    """
    result = {
        "strategy": "library",
        "format": "XLSX",
        "text": "",
        "tables": [],
        "images": [],
        "elapsed": 0.0,
    }

    start = time.time()

    try:
        from openpyxl import load_workbook

        wb = load_workbook(str(xlsx_path), data_only=True)
        sheet_texts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_str = " | ".join(str(v) if v is not None else "" for v in row)
                rows.append(row_str)

            sheet_text = f"## Sheet: {sheet_name}\n\n" + "\n".join(rows)
            sheet_texts.append(sheet_text)

            # 표 데이터 저장
            table_data = []
            for row in ws.iter_rows(values_only=True):
                table_data.append([str(v) if v is not None else "" for v in row])
            if table_data:
                result["tables"].append(table_data)

            # 차트 오브젝트 카운트
            if hasattr(ws, "_charts"):
                for chart in ws._charts:
                    result["images"].append(f"[chart: {chart.title or 'untitled'}]")

        result["text"] = "\n\n".join(sheet_texts)

    except ImportError:
        result["text"] = "[openpyxl 패키지 필요: pip install openpyxl]"

    result["elapsed"] = time.time() - start
    return result


def parse_xlsx_vllm(xlsx_path: Path) -> dict[str, Any]:
    """vLLM 기반 XLSX 파싱.

    시트를 테이블 이미지로 렌더링 후 LLaVA로 분석한다.

    Args:
        xlsx_path: XLSX 파일 경로.

    Returns:
        파싱 결과 딕셔너리.
    """
    result = {
        "strategy": "vllm",
        "format": "XLSX",
        "text": "",
        "tables": [],
        "images": [],
        "elapsed": 0.0,
    }

    start = time.time()

    try:
        from openpyxl import load_workbook

        wb = load_workbook(str(xlsx_path), data_only=True)
        sheet_texts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # 시트를 이미지로 렌더링
            img_path = _render_sheet_as_image(ws, sheet_name)
            if img_path and img_path.exists():
                caption = _call_llava(img_path)
                sheet_texts.append(
                    f"## Sheet: {sheet_name}\n\n{caption}"
                )
                result["images"].append(str(img_path))
                img_path.unlink(missing_ok=True)
            else:
                # 이미지 렌더링 실패 시 텍스트 대체
                rows = []
                for row in ws.iter_rows(values_only=True):
                    row_str = " | ".join(str(v) if v is not None else "" for v in row)
                    rows.append(row_str)
                sheet_texts.append(
                    f"## Sheet: {sheet_name}\n\n" + "\n".join(rows)
                )

        result["text"] = "\n\n".join(sheet_texts)

    except ImportError:
        result["text"] = "[openpyxl 패키지 필요: pip install openpyxl]"

    result["elapsed"] = time.time() - start
    return result


# ============================================================
# 유틸리티
# ============================================================

def _call_llava(image_path: Path) -> str:
    """LLaVA(Ollama)에 이미지를 전달하고 설명을 받는다.

    Args:
        image_path: 이미지 파일 경로.

    Returns:
        LLaVA 생성 캡션.
    """
    try:
        import base64
        import httpx

        with open(image_path, "rb") as f:
            img_b64 = base64.b64encode(f.read()).decode("utf-8")

        resp = httpx.post(
            f"{OLLAMA_BASE_URL}/api/chat",
            json={
                "model": VISION_MODEL,
                "messages": [
                    {
                        "role": "user",
                        "content": (
                            "Analyze this document image. "
                            "Extract all text, tables, and describe any charts or diagrams. "
                            "Output in structured Markdown format."
                        ),
                        "images": [img_b64],
                    }
                ],
                "stream": False,
            },
            timeout=120.0,
        )
        resp.raise_for_status()
        return resp.json()["message"]["content"]

    except ImportError:
        return "[httpx 패키지 필요: pip install httpx]"
    except Exception as e:
        return f"[LLaVA 분석 실패: {str(e)[:80]}]"


def _render_sheet_as_image(ws: Any, sheet_name: str) -> Path | None:
    """워크시트를 matplotlib 테이블 이미지로 렌더링한다.

    Args:
        ws: openpyxl 워크시트 객체.
        sheet_name: 시트 이름.

    Returns:
        생성된 이미지 경로 또는 None.
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        data = []
        for row in ws.iter_rows(values_only=True):
            data.append([str(v) if v is not None else "" for v in row])

        if not data:
            return None

        fig, ax = plt.subplots(figsize=(10, max(3, len(data) * 0.5)))
        ax.axis("off")

        table = ax.table(
            cellText=data[1:] if len(data) > 1 else data,
            colLabels=data[0] if len(data) > 1 else None,
            cellLoc="center",
            loc="center",
        )
        table.auto_set_font_size(False)
        table.set_fontsize(8)
        table.scale(1.2, 1.5)

        img_path = OUTPUTS_DIR / f"_vllm_sheet_{sheet_name}.png"
        img_path.parent.mkdir(parents=True, exist_ok=True)
        fig.savefig(str(img_path), dpi=150, bbox_inches="tight")
        plt.close(fig)
        return img_path

    except ImportError:
        return None


# ============================================================
# 비교 실험
# ============================================================

def compare_strategies(
    lib_result: dict[str, Any],
    vllm_result: dict[str, Any],
) -> None:
    """두 파싱 전략의 결과를 비교표로 출력한다.

    Args:
        lib_result: 라이브러리 파싱 결과.
        vllm_result: vLLM 파싱 결과.
    """
    table = Table(
        title=f"파싱 전략 비교 — {lib_result['format']}",
        show_lines=True,
    )
    table.add_column("항목", style="cyan", justify="center", width=14)
    table.add_column("라이브러리", style="yellow", width=20)
    table.add_column("vLLM (LLaVA)", style="green", width=20)

    table.add_row(
        "속도",
        f"{lib_result['elapsed']:.2f}초",
        f"{vllm_result['elapsed']:.2f}초",
    )
    table.add_row(
        "텍스트",
        f"{len(lib_result['text']):,}자",
        f"{len(vllm_result['text']):,}자",
    )
    table.add_row(
        "표 추출",
        f"{len(lib_result['tables'])}개",
        f"{len(vllm_result['tables'])}개" if vllm_result["tables"] else "MD 내 포함",
    )

    lib_img_desc = "추출 불가" if not lib_result["images"] else f"{len(lib_result['images'])}개 참조"
    vllm_img_desc = "설명 생성" if vllm_result["images"] else "없음"
    table.add_row("이미지", lib_img_desc, vllm_img_desc)

    has_chart_lib = any("chart" in str(i).lower() for i in lib_result["images"])
    has_chart_vllm = "chart" in vllm_result["text"].lower() or "그래프" in vllm_result["text"]
    table.add_row(
        "차트",
        "참조만" if has_chart_lib else "추출 불가",
        "내용 설명" if has_chart_vllm else "없음",
    )

    console.print(table)


def run_parser_comparison() -> None:
    """문서 파싱 전략 비교 데모를 실행한다."""
    console.rule("[bold blue]CH10 문서 파싱 전략 비교[/bold blue]")

    # --- PDF 비교 ---
    pdf_files = list(DATA_DIR.glob("*.pdf")) + list((DATA_DIR / "docs").rglob("*.pdf"))
    if pdf_files:
        pdf_path = pdf_files[0]
        console.print(f"\n[bold]PDF 비교: {pdf_path.name}[/bold]")

        lib_result = parse_pdf_library(pdf_path)
        console.print(f"  라이브러리 파싱 완료 ({lib_result['elapsed']:.2f}s)")

        vllm_result = parse_pdf_vllm(pdf_path)
        console.print(f"  vLLM 파싱 완료 ({vllm_result['elapsed']:.2f}s)")

        compare_strategies(lib_result, vllm_result)
    else:
        console.print("[yellow]PDF 파일이 없습니다. data/create_sample_docs.py를 먼저 실행하십시오.[/yellow]")
        _show_demo_comparison("PDF")

    # --- DOCX 비교 ---
    docx_files = list(DATA_DIR.glob("*.docx")) + list((DATA_DIR / "docs").rglob("*.docx"))
    if docx_files:
        docx_path = docx_files[0]
        console.print(f"\n[bold]DOCX 비교: {docx_path.name}[/bold]")

        lib_result = parse_docx_library(docx_path)
        console.print(f"  라이브러리 파싱 완료 ({lib_result['elapsed']:.2f}s)")

        vllm_result = parse_docx_vllm(docx_path)
        console.print(f"  vLLM 파싱 완료 ({vllm_result['elapsed']:.2f}s)")

        compare_strategies(lib_result, vllm_result)
    else:
        console.print("[yellow]DOCX 파일이 없습니다.[/yellow]")
        _show_demo_comparison("DOCX")

    # --- XLSX 비교 ---
    xlsx_files = list(DATA_DIR.glob("*.xlsx")) + list((DATA_DIR / "docs").rglob("*.xlsx"))
    if xlsx_files:
        xlsx_path = xlsx_files[0]
        console.print(f"\n[bold]XLSX 비교: {xlsx_path.name}[/bold]")

        lib_result = parse_xlsx_library(xlsx_path)
        console.print(f"  라이브러리 파싱 완료 ({lib_result['elapsed']:.2f}s)")

        vllm_result = parse_xlsx_vllm(xlsx_path)
        console.print(f"  vLLM 파싱 완료 ({vllm_result['elapsed']:.2f}s)")

        compare_strategies(lib_result, vllm_result)
    else:
        console.print("[yellow]XLSX 파일이 없습니다.[/yellow]")
        _show_demo_comparison("XLSX")

    # --- 결론 ---
    console.rule("[bold green]비교 완료[/bold green]")
    console.print(
        "\n[bold]전략 선택 가이드:[/bold]\n"
        "  - 텍스트 위주 문서 → 라이브러리 파싱 (빠르고 정확)\n"
        "  - 이미지/차트 포함 문서 → vLLM 파싱 (시각 정보 활용)\n"
        "  - 표 추출이 중요 → 라이브러리(pdfplumber) 권장\n"
        "  - 스캔 PDF → vLLM(LLaVA) 필수\n"
    )


def _show_demo_comparison(fmt: str) -> None:
    """파일이 없을 때 데모 비교표를 출력한다."""
    table = Table(title=f"파싱 전략 비교 — {fmt} (데모)", show_lines=True)
    table.add_column("항목", style="cyan", justify="center", width=14)
    table.add_column("라이브러리", style="yellow", width=20)
    table.add_column("vLLM (LLaVA)", style="green", width=20)

    if fmt == "PDF":
        table.add_row("속도", "0.3초", "4.2초")
        table.add_row("텍스트", "1,240자", "1,180자")
        table.add_row("표 추출", "2개 (깨짐 가능)", "2개 (정확)")
        table.add_row("이미지", "추출 불가", "설명 생성")
        table.add_row("차트", "추출 불가", "내용 설명")
    elif fmt == "DOCX":
        table.add_row("속도", "0.1초", "3.8초")
        table.add_row("텍스트", "890자", "920자")
        table.add_row("표 추출", "1개 (정확)", "MD 내 포함")
        table.add_row("이미지", "파일 참조", "설명 생성")
        table.add_row("차트", "참조만", "내용 설명")
    else:  # XLSX
        table.add_row("속도", "0.2초", "5.1초")
        table.add_row("텍스트", "650자", "580자")
        table.add_row("표 추출", "3개 (정확)", "MD 내 포함")
        table.add_row("이미지", "추출 불가", "없음")
        table.add_row("차트", "오브젝트 참조", "내용 설명")

    console.print(table)


if __name__ == "__main__":
    # 프로젝트 루트를 sys.path에 추가
    _PROJECT_ROOT = str(BASE_DIR)
    if _PROJECT_ROOT not in sys.path:
        sys.path.insert(0, _PROJECT_ROOT)

    run_parser_comparison()
